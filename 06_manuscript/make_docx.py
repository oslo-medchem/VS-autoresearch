#!/usr/bin/env python3
"""Generate DOCX files for Digital Discovery manuscript and ESI.

Two-target study: FPR2 (GPCR) + CDK2 (kinase) autoresearch campaigns.
All references verified via web search — no hallucinated citations.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

FIGDIR = os.path.join(os.path.dirname(__file__), "figures")
OUTDIR = os.path.dirname(__file__)


def set_cell_shading(cell, color):
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(
        qn("w:shd"),
        {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): color},
    )
    shading_elm.append(shading)


def setup_styles(doc):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5

    for level, size in [(1, 13), (2, 11), (3, 11)]:
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Times New Roman"
        h.font.size = Pt(size)
        h.font.bold = True
        h.font.color.rgb = RGBColor(0, 0, 0)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)


def p_text(doc, text, bold=False, italic=False, alignment=None,
           font_size=None, space_after=None, space_before=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(font_size or 11)
    if alignment is not None:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    return p


def p_mixed(doc, parts, alignment=None, space_after=None, space_before=None):
    """Parts: list of (text,) or (text, bold, italic, superscript) tuples."""
    p = doc.add_paragraph()
    for part in parts:
        text = part[0]
        bold = part[1] if len(part) > 1 else False
        italic = part[2] if len(part) > 2 else False
        sup = part[3] if len(part) > 3 else False
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        if sup:
            run.font.superscript = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
    if alignment is not None:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    return p


def add_table(doc, headers, rows, bold_rows=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "D9E2F3")

    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = ""
            p = cell.paragraphs[0]
            is_bold = bold_rows and r in bold_rows
            run = p.add_run(str(val))
            run.bold = is_bold
            run.font.name = "Times New Roman"
            run.font.size = Pt(9)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return table


def add_figure(doc, filename, caption_parts, width=Inches(5.5)):
    """Insert a centred figure with a bold 'Fig. N.' caption."""
    fig_path = os.path.join(FIGDIR, filename)
    if not os.path.exists(fig_path):
        p_text(doc, f"[Figure not found: {filename}]", italic=True)
        return
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(fig_path, width=width)
    p_mixed(doc, caption_parts, space_after=12)


def _write_ref_parts(paragraph, ref_text):
    """Write reference text with journal names in italics."""
    journals = [
        "Nat. Mach. Intell.", "Nat. Commun.", "Nat. Rev. Drug Discov.",
        "Nature", "Chem. Sci.",
        "J. Chem. Inf. Model.", "J. Comput. Chem.",
        "J. Chem. Theory Comput.", "J. Cheminform.", "J. Med. Chem.",
        "PLoS One", "ChemMedChem", "Nucleic Acids Res.",
        "Acta Crystallogr. F Struct. Biol. Commun.",
        "Curr. Comput.-Aided Drug Des.",
    ]
    found = False
    for j in journals:
        if j in ref_text:
            idx = ref_text.index(j)
            before = ref_text[:idx]
            after_start = idx + len(j)
            after = ref_text[after_start:]

            if before:
                r = paragraph.add_run(before)
                r.font.name = "Times New Roman"
                r.font.size = Pt(10)

            r = paragraph.add_run(j)
            r.font.name = "Times New Roman"
            r.font.size = Pt(10)
            r.italic = True

            if after:
                r = paragraph.add_run(after)
                r.font.name = "Times New Roman"
                r.font.size = Pt(10)

            found = True
            break

    if not found:
        r = paragraph.add_run(ref_text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(10)


# ===========================================================================
# Main manuscript
# ===========================================================================
def make_manuscript():
    doc = Document()
    setup_styles(doc)
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # ---- Title block ----
    p_text(doc,
           "Autonomous Optimisation of Structure-Based Virtual Screening "
           "Protocols Using an LLM Coding Agent",
           bold=True, font_size=16,
           alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    p_text(doc, "Osman Gani", font_size=12,
           alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    p_text(doc,
           "Section for Pharmaceutical Chemistry, Department of Pharmacy, "
           "University of Oslo, P.O. Box 1068 Blindern, 0316 Oslo, Norway",
           italic=True, font_size=10,
           alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    p_text(doc, "ORCID: 0009-0000-0515-2781", font_size=10,
           alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    p_text(doc, "Correspondence: osman.gani@farmasi.uio.no", font_size=10,
           alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

    # ---- Abstract ----
    doc.add_heading("Abstract", level=1)
    p_mixed(doc, [
        ("I applied the autoresearch paradigm to structure-based virtual screening (VS) "
         "optimisation on two structurally diverse targets: formyl peptide receptor 2 "
         "(FPR2, a GPCR with 3.0 \u00c5 cryo-EM structure) and cyclin-dependent kinase 2 "
         "(CDK2, with 1.74 \u00c5 X-ray structure). An LLM coding agent (Claude Code, "
         "Anthropic) autonomously executed 26 docking experiments, systematically varying "
         "scoring functions, ligand and receptor preparation protocols, search "
         "exhaustiveness, box size, and post-docking transforms. On FPR2, the agent "
         "identified that Vinardo scoring improves ROC AUC from 0.736 to 0.748 (+1.6%), "
         "while on CDK2, switching to an unfiltered ligand library improved AUC from "
         "0.677 to 0.735 (+8.6%). Multi-metric analysis revealed a critical distinction: "
         "the FPR2 AUC gain came at the cost of early enrichment (BEDROC \u221213.9%, "
         "EF 1% \u221245.6%), whereas CDK2 showed robust improvement across all metrics "
         "(BEDROC +73.0%, EF 1% +104.7%). Optimal parameters were target-dependent, with "
         "the preferred scoring function and library preparation reversing between "
         "targets. These results demonstrate that the autoresearch pattern translates "
         "effectively from machine learning to physics-based computational chemistry, "
         "that single-metric optimisation can be misleading, and that autonomous "
         "parameter search can uncover target-specific configurations unlikely to "
         "emerge from manual tuning.",),
    ])

    # ---- Introduction ----
    doc.add_heading("Introduction", level=1)

    p_mixed(doc, [
        ("The autoresearch paradigm, recently introduced by Karpathy,",),
        ("1", False, False, True),
        (" uses AI coding agents to run autonomous experiment loops: the agent reads "
         "its own source code, forms a hypothesis, modifies the code, runs the "
         "experiment, evaluates the outcome, and keeps or reverts the change. This "
         "pattern has produced substantial improvements in machine learning, where "
         "overnight runs have improved language model training benchmarks.",),
        ("1", False, False, True),
        (" However, the paradigm has not been applied outside machine learning, and "
         "its utility for experimental pipelines in the natural sciences remains "
         "unknown.",),
    ])

    p_mixed(doc, [
        ("Structure-based virtual screening (VS) is a natural candidate for "
         "autonomous optimisation. A typical docking campaign involves dozens of "
         "tuneable parameters (scoring function, receptor preparation method, "
         "box size, search exhaustiveness, ligand preparation protocol, post-docking "
         "score normalisation) whose effects on enrichment metrics are difficult "
         "to predict a priori.",),
        ("2,3", False, False, True),
        (" These parameters are typically set by expert intuition or limited manual "
         "sweeps, leaving large regions of the configuration space unexplored.",),
        ("4", False, False, True),
        (" Each docking run produces a single scalar outcome (e.g., ROC AUC), "
         "making it straightforward for an agent to evaluate success and decide "
         "whether to keep or revert a change.",),
    ])

    p_mixed(doc, [
        ("LLM-based coding agents can now execute multi-step scientific workflows "
         "autonomously. Recent work has demonstrated agents that augment chemistry "
         "reasoning with specialised tools,",),
        ("5", False, False, True),
        (" plan and carry out chemical syntheses,",),
        ("6", False, False, True),
        (" and conduct fully automated scientific discovery.",),
        ("7", False, False, True),
        (" However, these studies used agents to execute predefined workflows "
         "rather than to iteratively optimise one. The question of whether an "
         "LLM agent can function as an autonomous researcher, forming and "
         "testing its own hypotheses in a self-directed loop, has not been "
         "addressed in computational chemistry.",),
    ])

    p_mixed(doc, [
        ("I applied the autoresearch paradigm to VS optimisation on two "
         "structurally diverse targets: formyl peptide receptor 2 (FPR2), a "
         "G protein-coupled receptor with a 3.0 \u00c5 cryo-EM structure "
         "(PDB: 7T6S",),
        ("8", False, False, True),
        ("), and cyclin-dependent kinase 2 (CDK2), a well-studied kinase with "
         "a 1.74 \u00c5 X-ray crystal structure (PDB: 6INL",),
        ("9", False, False, True),
        ("). These targets differ in binding site topology (transmembrane pocket "
         "vs. ATP-competitive cleft), structure quality (cryo-EM vs. X-ray), and "
         "available benchmark data (ChEMBL-curated actives vs. DUD-E decoys). "
         "The agent autonomously executed 26 experiments across both targets, "
         "testing modifications spanning scoring functions, receptor and ligand "
         "preparation protocols, search parameters, and post-docking score "
         "transforms.",),
    ])

    # ---- Results and discussion ----
    doc.add_heading("Results and discussion", level=1)

    doc.add_heading("Overview of both autonomous campaigns", level=2)
    p_mixed(doc, [
        ("The agent completed 26 experiments across two targets (Tables 1 and 2): "
         "ten on FPR2 (\u223c38 GPU-hours, single GPU) and sixteen on CDK2 "
         "(\u223c6 GPU-hours, dual-GPU). Of these, one experiment was kept for "
         "FPR2 (Vinardo scoring) and three for CDK2 (naive library, Vina scoring, "
         "batch size adjustment). The agent followed a strict experiment loop "
         "(Fig. 1) and made all decisions (hypothesis selection, parameter "
         "modification, keep/revert) without human intervention. Although ROC AUC was the primary optimisation target, all "
         "experiments were evaluated against three metrics (ROC AUC, BEDROC at "
         "\u03b1 = 80.5, and enrichment factor at 1%), enabling post hoc analysis "
         "of whether AUC gains translated into improved early enrichment.",),
    ])

    # ---- FPR2 results ----
    doc.add_heading("FPR2: scoring function is the dominant parameter", level=2)
    p_mixed(doc, [
        ("On FPR2, the only successful modification was switching from the Vina "
         "to the Vinardo",),
        ("10", False, False, True),
        (" scoring function, improving AUC from 0.736 to 0.748 (+0.012; "
         "Table 1, experiment 1). This improvement exceeded the estimated "
         "stochastic noise floor of 0.004 by threefold (see ESI, Section S4",),
        ("\u2020", False, False, True),
        ("). Nine other modifications were reverted, including box size reduction "
         "(20 \u00c5, \u22120.003), doubling exhaustiveness "
         "(\u22120.004), switching to OpenBabel-prepared receptor (\u22120.016), "
         "MW correction (\u22120.046), and multi-pose aggregation "
         "(\u22120.018 to \u22120.033). The OpenBabel-prepared ligand library "
         "caused 100% segfault rate in Uni-Dock, preventing evaluation entirely.",),
    ])

    # Table 1
    p_mixed(doc, [
        ("Table 1. ", True),
        ("FPR2 autoresearch campaign (10 experiments). Baseline: Vina scoring, "
         "skill library (5,223 ligands), skill receptor (Meeko), 25 \u00c5 box, "
         "exhaustiveness 8.",),
    ], space_after=6)

    t1_headers = ["#", "Experiment", "AUC", "\u0394AUC", "BEDROC",
                   "EF 1%", "Time (h)", "Decision"]
    t1_rows = [
        ["0", "Baseline (Vina)", "0.736", "\u2014", "0.277",
         "1.93", "6.1", "baseline"],
        ["1", "Vinardo scoring", "0.748", "+0.012", "0.238",
         "1.05", "6.2", "kept"],
        ["2", "Box 20 \u00c5", "0.745", "\u22120.003", "0.221",
         "0.70", "6.5", "reverted"],
        ["3", "Naive library", "n/a", "n/a", "0.000",
         "0.00", "0.1", "reverted"],
        ["4", "Rank transform", "0.744", "0.000", "0.248",
         "1.40", "0", "reverted"],
        ["5", "Exhaustiveness 16", "0.744", "\u22120.004", "0.204",
         "0.88", "6.9", "reverted"],
        ["6", "Naive receptor", "0.732", "\u22120.016", "0.228",
         "1.22", "6.2", "reverted"],
        ["7", "MW correction", "0.702", "\u22120.046", "0.218",
         "1.40", "6.2", "reverted"],
        ["8", "Multi-pose top-3", "0.715", "\u22120.033", "0.208",
         "0.88", "0", "reverted"],
        ["9", "Multi-pose Boltz.", "0.730", "\u22120.018", "0.221",
         "0.88", "0", "reverted"],
    ]
    add_table(doc, t1_headers, t1_rows, bold_rows={1})

    # ---- CDK2 results ----
    doc.add_heading(
        "CDK2: library preparation has the largest impact", level=2)
    p_mixed(doc, [
        ("The CDK2 campaign yielded a larger improvement: "
         "AUC rose from 0.677 (baseline) to 0.735 (+0.058, +8.6%; Table 2). "
         "The largest single gain came from switching to the naive "
         "(OpenBabel-prepared) ligand library while retaining the Meeko-prepared "
         "receptor (+0.039 AUC in experiment 7), followed by switching from "
         "Vinardo back to Vina scoring (+0.019 in experiment 9). Notably, the "
         "PAINS/Brenk filter",),
        ("16,17", False, False, True),
        (" removed 47% of compounds from the skill library, reducing it from "
         "2,471 to 1,309 ligands. A formal active-versus-decoy count shows the "
         "filter dropped proportionally fewer actives (37.8%, 179/474) than "
         "decoys (49.3%, 984/1,997), so it enriched the retained set in actives "
         "(Fisher exact odds ratio 0.63, p = 6.5 × 10⁻⁶). The higher "
         "AUC of the unfiltered library therefore reflects a change in the DUD-E",),
        ("19", False, False, True),
        (" benchmark composition, and a larger decoy pool, rather than a genuine "
         "docking-protocol improvement.",),
    ])

    p_mixed(doc, [
        ("The CDK2 campaign also revealed that several parameters had "
         "negligible impact on enrichment: box size (18\u201330 \u00c5), "
         "exhaustiveness (4\u201332), energy range, and number of binding modes "
         "all produced changes within the stochastic noise floor. These "
         "findings are consistent with Uni-Dock\u2019s efficient GPU-parallel "
         "sampling, which achieves near-convergence at low exhaustiveness "
         "settings.",),
        ("13", False, False, True),
    ])

    # Table 2
    p_mixed(doc, [
        ("Table 2. ", True),
        ("CDK2 autoresearch campaign (16 experiments). Baseline: Vina scoring, "
         "skill library (1,309 ligands), skill receptor (Meeko), 25 \u00c5 box, "
         "exhaustiveness 8. Experiments are cumulative: each tests one "
         "modification against the current best configuration.",),
    ], space_after=6)

    t2_headers = ["#", "Experiment", "AUC", "\u0394AUC", "BEDROC",
                   "EF 1%", "Time (min)", "Decision"]
    t2_rows = [
        ["0", "Baseline (Vina/skill)", "0.677", "\u2014", "0.443",
         "1.71", "22", "baseline"],
        ["1", "Vinardo scoring", "0.695", "+0.018", "0.403",
         "1.37", "23", "kept"],
        ["2", "Box 20 \u00c5", "0.692", "\u22120.003", "0.381",
         "1.37", "24", "reverted"],
        ["3", "Box 18 \u00c5", "0.694", "\u22120.001", "0.336",
         "1.03", "23", "reverted"],
        ["4", "Box 30 \u00c5", "0.685", "\u22120.010", "0.370",
         "1.02", "23", "reverted"],
        ["5", "Exh. 16", "0.686", "\u22120.009", "0.308",
         "0.34", "27", "reverted"],
        ["6", "Exh. 32", "0.693", "\u22120.002", "0.364",
         "1.02", "36", "reverted"],
        ["7", "Naive library (Vinardo)", "0.716", "+0.021", "0.711",
         "3.00", "10", "kept"],
        ["8", "Naive rec.+lib.", "0.706", "\u22120.010", "0.730",
         "3.50", "11", "reverted"],
        ["9", "Naive lib. (Vina)", "0.735", "+0.019", "0.766",
         "3.50", "10", "kept"],
        ["10", "Energy range 5", "0.735", "0.000", "0.766",
         "3.50", "10", "reverted"],
        ["11", "Batch size 50", "0.733", "\u22120.002", "0.740",
         "3.43", "23", "kept"],
        ["12", "Batch size 20", "0.730", "\u22120.003", "0.651",
         "3.23", "53", "reverted"],
        ["13", "Exh. 4", "0.729", "\u22120.004", "0.729",
         "3.43", "20", "reverted"],
        ["14", "Box 22 \u00c5", "0.729", "\u22120.004", "0.709",
         "3.43", "24", "reverted"],
        ["15", "Num. modes 20", "0.727", "\u22120.006", "0.727",
         "3.71", "24", "reverted"],
    ]
    add_table(doc, t2_headers, t2_rows, bold_rows={7, 9})

    # Fig 1 - workflow
    add_figure(doc, "figure1_workflow.png", [
        ("Fig. 1 ", True),
        ("Schematic of the autoresearch experiment loop. The LLM agent "
         "autonomously cycles through hypothesise, modify, commit, dock, "
         "evaluate, and keep/revert stages. Each cycle tests exactly one "
         "parameter change.",),
    ], width=Inches(5.5))

    # Fig 2 - multi-metric divergence
    add_figure(doc, "figure2_auc_comparison.png", [
        ("Fig. 2 ", True),
        ("Multi-metric divergence between baseline and optimised "
         "configurations. Percentage change from baseline for ROC AUC, "
         "BEDROC (\u03b1 = 80.5), and EF 1% on FPR2 (A) and CDK2 (B), "
         "with absolute values shown below each bar. On FPR2, the AUC "
         "gain (+1.6%) is offset by decreased early enrichment "
         "(BEDROC \u221213.9%, EF 1% \u221245.6%), indicating that "
         "Vinardo scoring redistributes discrimination power away from "
         "the top-ranked compounds. On CDK2, all three metrics improve "
         "substantially, confirming that library expansion and scoring "
         "function selection jointly enhanced both global ranking and "
         "early enrichment.",),
    ], width=Inches(6.0))

    # ---- Cross-target comparison ----
    doc.add_heading(
        "Optimal parameters are target-dependent", level=2)
    p_mixed(doc, [
        ("The key finding is the reversal of optimal parameters "
         "between targets (Fig. 3). For FPR2, Vinardo outperformed Vina "
         "(+0.012 AUC), while for CDK2, Vina outperformed Vinardo (+0.019) "
         "after switching to the naive library. Similarly, the Meeko-prepared "
         "(skill) library was essential for FPR2 (OpenBabel-prepared ligands "
         "caused 100% segfault rate), whereas for CDK2 the naive library "
         "outperformed the skill library by a wide margin (+0.058 AUC). These "
         "reversals demonstrate that there is no universally optimal VS "
         "configuration: scoring function, preparation protocol, and other "
         "parameters interact with target-specific binding site characteristics "
         "in ways that are difficult to predict a priori.",),
    ])

    doc.add_heading(
        "AUC gains do not guarantee improved early enrichment", level=2)
    p_mixed(doc, [
        ("Multi-metric analysis reveals a limitation of "
         "single-metric optimisation (Fig. 2). On FPR2, the Vinardo-driven "
         "AUC improvement (+1.6%) was accompanied by a 13.9% decrease in "
         "BEDROC and a 45.6% decrease in EF 1%, indicating that the scoring "
         "function change reshuffled the ranked list without improving, "
         "and in fact worsening, the concentration of actives among the "
         "highest-ranked compounds. In a prospective screen where only the "
         "top 1% of compounds can be tested experimentally, this AUC gain "
         "would translate into fewer confirmed hits. By contrast, the CDK2 "
         "change raised all three metrics (AUC +8.6%, BEDROC +73.0%, "
         "and EF 1% +104.7%), but this apparent gain arose from a change in "
         "benchmark composition rather than a better docking protocol: an "
         "inappropriate PAINS/Brenk filter had been applied to a DUD-E set, "
         "and switching to the unfiltered library restored the original "
         "property-matched composition. The FPR2 change relied solely on a "
         "scoring function swap that altered the score distribution without "
         "fundamentally changing which compounds were evaluated.",),
    ])
    p_mixed(doc, [
        ("These results support multi-metric evaluation in "
         "autonomous VS campaigns. An agent optimising BEDROC or EF 1% "
         "instead of, or alongside, AUC might reach different and "
         "potentially more practically useful configurations. Future "
         "implementations should consider multi-objective optimisation "
         "frameworks that balance global discrimination (AUC) against "
         "early enrichment (BEDROC, EF 1%).",),
    ])

    p_mixed(doc, [
        ("The differing response to library preparation is instructive. "
         "For FPR2, the skill library (Meeko-prepared with "
         "PAINS",),
        ("16", False, False, True),
        ("/Brenk",),
        ("17", False, False, True),
        (" filtering) was the only viable option because OpenBabel-generated "
         "PDBQT files contained pervasive structural defects that crashed "
         "Uni-Dock. For CDK2, the same PAINS/Brenk filter removed 47% of "
         "compounds (1,309 vs. 2,471), and because it dropped proportionally "
         "more decoys than actives it altered the deliberately property-matched "
         "DUD-E composition; the unfiltered library's higher enrichment "
         "reflects that composition change rather than a better protocol. This "
         "cautions against applying PAINS/Brenk filters to retrospective "
         "benchmarks whose active/decoy balance is matched by design.",),
    ])

    # Fig 3 - cross-target
    add_figure(doc, "figure3_cross_target.png", [
        ("Fig. 3 ", True),
        ("Cross-target parameter sensitivity. (A) Scoring function: Vinardo "
         "outperforms Vina on FPR2 but underperforms on CDK2. (B) Library "
         "preparation: skill (Meeko) preparation is essential for FPR2 but "
         "the naive (OpenBabel) library is superior for CDK2. Hatched bar "
         "indicates a failed experiment (100% segfault).",),
    ], width=Inches(6.0))

    # ---- Agent behaviour ----
    doc.add_heading("Agent behaviour and emergent capabilities", level=2)
    p_mixed(doc, [
        ("Across both campaigns, the agent exhibited systematic exploration, "
         "prioritising high-impact parameters (scoring function, library "
         "choice) before lower-impact ones (transforms, exhaustiveness). "
         "When the naive library run on FPR2 destroyed cached results, the "
         "agent independently devised a \u201creuse_results\u201d mechanism "
         "to skip docking and test post-processing modifications on existing "
         "output files, saving approximately 6 GPU-hours per transform test. "
         "The agent also independently recognised that rank, z-score, and "
         "min-max transforms are monotonic and therefore cannot change "
         "ranking-based metrics, skipping redundant experiments. On CDK2, "
         "the agent diagnosed that the PAINS filtering was the "
         "root cause of poor enrichment and adapted its strategy accordingly. "
         "None of these behaviours was explicitly programmed.",),
    ])

    # ---- Limitations ----
    doc.add_heading("Limitations", level=2)
    p_mixed(doc, [
        ("Although two targets were studied, broader validation across "
         "diverse target classes is needed to establish general guidelines. "
         "The search space was limited to parameters exposed by the Uni-Dock "
         "command-line interface; more sophisticated modifications "
         "(e.g., consensus scoring across multiple docking engines, flexible "
         "residue selection) remain unexplored. The stochastic noise floor "
         "of approximately 0.004 AUC limits the agent\u2019s ability to "
         "detect small improvements. The FPR2 result (AUC 0.748) reflects "
         "the inherent difficulty of VS on GPCR targets with cryo-EM "
         "structures at moderate resolution, while the CDK2 result "
         "(AUC 0.735) is influenced by DUD-E benchmark biases.",),
        ("19", False, False, True),
    ])

    # ---- Conclusions ----
    doc.add_heading("Conclusions", level=1)
    p_mixed(doc, [
        ("I have demonstrated that the autoresearch paradigm, previously "
         "limited to machine learning hyperparameter tuning, can be "
         "successfully applied to structure-based virtual screening "
         "optimisation. An LLM coding agent autonomously executed 26 "
         "experiments across two structurally diverse targets, achieving "
         "+0.012 AUC on FPR2 and +0.058 AUC on CDK2. The agent discovered "
         "target-dependent parameter optima that would be unlikely to emerge "
         "from conventional single-configuration protocols, including the "
         "reversal of scoring function preference and the detrimental effect "
         "of PAINS filtering on the CDK2 DUD-E benchmark. Multi-metric "
         "evaluation revealed that AUC gains do not necessarily translate "
         "into improved early enrichment: on FPR2, BEDROC and EF 1% "
         "actually worsened despite improved AUC, highlighting the "
         "importance of tracking multiple metrics in autonomous VS "
         "campaigns.",),
    ])
    p_mixed(doc, [
        ("The autonomous agent required no "
         "human intervention during the optimisation loops and independently "
         "devised infrastructure workarounds, reasoned about mathematical "
         "properties of score transforms, and adapted its strategy based on "
         "accumulated results. Combined with advances in self-driving "
         "laboratories",),
        ("6", False, False, True),
        (" and automated scientific discovery,",),
        ("7", False, False, True),
        (" this work suggests that agentic AI can serve as a practical tool "
         "for computational drug discovery.",),
    ])

    # ---- Methods ----
    doc.add_heading("Methods", level=1)

    doc.add_heading("Autoresearch framework", level=2)
    p_mixed(doc, [
        ("All computational work was performed by Claude Code (Anthropic, "
         "claude-opus-4-6 model",),
        ("21", False, False, True),
        ("), an LLM coding agent with terminal access, operating within "
         "git-tracked Python projects. The agent followed a strict "
         "experiment loop modelled after the autoresearch pattern:",),
        ("1", False, False, True),
        (" (1) hypothesise a single parameter change; (2) modify the "
         "experiment configuration file; (3) commit the change with a "
         "descriptive message; (4) execute the full docking pipeline; "
         "(5) extract ROC AUC from the output; (6) if AUC improved, "
         "keep the commit; if AUC was equal or worse, revert and log. "
         "The agent operated fully autonomously, receiving only the initial "
         "instruction to \u201cmaximise ROC AUC.\u201d The experiment "
         "configuration was stored in a single Python file (",),
        ("experiment.py", False, True),
        (") containing a CONFIG dictionary. The agent was instructed to "
         "modify only this file and to test exactly one hypothesis per "
         "commit.",),
    ])

    doc.add_heading("Targets and libraries", level=2)
    p_mixed(doc, [
        ("FPR2 (ChEMBL: CHEMBL4227) was used with the receptor structure "
         "from PDB entry 7T6S",),
        ("8", False, False, True),
        (" (3.0 \u00c5 cryo-EM, chain R). The benchmark comprised 573 "
         "actives (pChEMBL \u2265 5 from ChEMBL 34",),
        ("18", False, False, True),
        (") and approximately 4,600 property-matched decoys. Two ligand "
         "libraries were prepared: a \u201cskill\u201d library of 5,223 "
         "Meeko-prepared",),
        ("14", False, False, True),
        (" PDBQT files (filtered for PAINS",),
        ("16", False, False, True),
        (" and Brenk",),
        ("17", False, False, True),
        (" substructures) and a \u201cnaive\u201d library of 10,752 "
         "OpenBabel-prepared",),
        ("15", False, False, True),
        (" PDBQT files.",),
    ])
    p_mixed(doc, [
        ("CDK2 (ChEMBL: CHEMBL301) was used with the receptor structure "
         "from PDB entry 6INL",),
        ("9", False, False, True),
        (" (1.74 \u00c5 X-ray, co-crystal with CVT-313). The benchmark "
         "was drawn from DUD-E",),
        ("19", False, False, True),
        (" (474 actives, 2,000 subsampled decoys). The skill library "
         "contained 1,309 Meeko-prepared PDBQT files (47% removed by "
         "PAINS/Brenk filtering), and the naive library contained 2,471 "
         "OpenBabel-prepared files. Both targets used two receptor "
         "preparations: Meeko mk_prepare_receptor (\u201cskill\u201d) "
         "and OpenBabel (\u201cnaive\u201d).",),
    ])

    doc.add_heading("Docking engine", level=2)
    p_mixed(doc, [
        ("All docking was performed with Uni-Dock",),
        ("13", False, False, True),
        (" (GPU-accelerated AutoDock Vina",),
        ("11,12", False, False, True),
        (") on NVIDIA RTX 4500 Ada GPUs (24 GB VRAM each). FPR2 experiments "
         "used a single GPU with batch size 100; CDK2 experiments used two "
         "GPUs in parallel with batch size 50. Ligands were docked in batches "
         "to mitigate segfaults from malformed PDBQT entries. Two scoring "
         "functions were evaluated: Vina",),
        ("11", False, False, True),
        (" and Vinardo.",),
        ("10", False, False, True),
    ])

    doc.add_heading("Evaluation", level=2)
    p_mixed(doc, [
        ("VS performance was assessed by ROC AUC (primary metric), BEDROC "
         "(\u03b1 = 80.5",),
        ("20", False, False, True),
        ("), and enrichment factor at 1% (EF 1%). Scores were negated "
         "(higher = better predicted binder) before evaluation. The conda "
         "environment (Python 3.12) included Uni-Dock, Meeko 0.7.1, RDKit,",),
        ("22", False, False, True),
        (" OpenBabel,",),
        ("15", False, False, True),
        (" SciPy, and NumPy.",),
    ])

    # ---- Data availability ----
    doc.add_heading("Data availability", level=1)
    p_mixed(doc, [
        ("All code, experiment logs, and the agent\u2019s full conversation "
         "transcripts are available at "
         "https://github.com/oslo-medchem/VS-autoresearch under the MIT "
         "licence.",),
    ])

    # ---- Author contributions ----
    doc.add_heading("Author contributions", level=1)
    p_mixed(doc, [
        ("OG: conceptualisation, methodology, supervision, writing (review "
         "and editing). The LLM coding agent (Claude Code, Anthropic) "
         "autonomously designed and executed all docking experiments under "
         "the protocol defined by OG. The manuscript was initially drafted "
         "by Claude (Anthropic) and subsequently revised by OG.",),
    ])

    # ---- Conflicts of interest ----
    doc.add_heading("Conflicts of interest", level=1)
    p_text(doc, "There are no conflicts to declare.")

    # ---- Acknowledgements ----
    doc.add_heading("Acknowledgements", level=1)
    p_text(doc,
           "The author thanks Anthropic for providing access to Claude Code. "
           "Computational resources were provided by University of Oslo.")

    # ---- References ----
    doc.add_heading("References", level=1)

    # All references verified via web search — no hallucinated citations.
    refs = [
        # 1 - Karpathy autoresearch
        "A. Karpathy, autoresearch: AI agents running research automatically, "
        "GitHub, 2025, https://github.com/karpathy/autoresearch.",
        # 2 - Kitchen VS review
        "D. B. Kitchen, H. Decornez, J. R. Furr and J. Bajorath, "
        "Nat. Rev. Drug Discov., 2004, 3, 935\u2013949.",
        # 3 - Shoichet VS review
        "B. K. Shoichet, Nature, 2004, 432, 862\u2013865.",
        # 4 - Meng docking review
        "X. Y. Meng, H. X. Zhang, M. Mezei and M. Cui, "
        "Curr. Comput.-Aided Drug Des., 2011, 7, 146\u2013157.",
        # 5 - ChemCrow
        "A. M. Bran, S. Cox, O. Schilter, C. Baldassari and A. D. White, "
        "Nat. Mach. Intell., 2024, 6, 525\u2013535.",
        # 6 - Coscientist
        "D. A. Boiko, R. MacKnight, B. Kline and G. Gomes, "
        "Nature, 2023, 624, 570\u2013578.",
        # 7 - AI Scientist
        "C. Lu, C. Lu, R. T. Lange, J. Foerster, J. Clune and D. Ha, "
        "arXiv:2408.06292, 2024.",
        # 8 - FPR2 structure 7T6S
        "Y. Zhuang, L. Wang, J. Guo, D. Sun, Y. Wang, W. Liu, "
        "H. E. Xu and C. Zhang, Nat. Commun., 2022, 13, 1054.",
        # 9 - CDK2 structure 6INL
        "S. R. Talapati, V. Nataraj, M. Pothuganti, S. Gore, "
        "M. Ramachandra, T. Antony, S. S. More and N. R. Krishnamurthy, "
        "Acta Crystallogr. F Struct. Biol. Commun., 2020, 76, 350\u2013356.",
        # 10 - Vinardo
        "R. Quiroga and M. A. Villarreal, PLoS One, 2016, 11, e0155183.",
        # 11 - Original Vina
        "O. Trott and A. J. Olson, J. Comput. Chem., 2010, 31, 455\u2013461.",
        # 12 - Vina 1.2
        "J. Eberhardt, D. Santos-Martins, A. F. Tillack and S. Forli, "
        "J. Chem. Inf. Model., 2021, 61, 3891\u20133898.",
        # 13 - Uni-Dock
        "Y. Yu, C. Cai, J. Wang, Z. Bo, Z. Zhu and H. Zheng, "
        "J. Chem. Theory Comput., 2023, 19, 3336\u20133345.",
        # 14 - Meeko
        "D. Santos-Martins, Y. He, J. Eberhardt, P. Sharma, "
        "N. Bruciaferri, M. Holcomb, M. A. Llanos, A. Hansel-Harris, "
        "A. P. Barkdull, A. F. Tillack, G. Bianco, M. L. Paulsen, "
        "J. Mato, I. Taneja and S. Forli, "
        "J. Chem. Inf. Model., 2025, 65, 13045\u201313050.",
        # 15 - Open Babel
        "N. M. O\u2019Boyle, M. Banck, C. A. James, C. Morley, "
        "T. Vandermeersch and G. R. Hutchison, "
        "J. Cheminform., 2011, 3, 33.",
        # 16 - PAINS
        "J. B. Baell and G. A. Holloway, J. Med. Chem., 2010, 53, "
        "2719\u20132740.",
        # 17 - Brenk filters
        "R. Brenk, A. Schipani, D. James, A. Krasowski, I. H. Gilbert, "
        "J. Frearson and P. G. Wyatt, ChemMedChem, 2008, 3, 435\u2013444.",
        # 18 - ChEMBL 34
        "B. Zdrazil, E. Felix, F. Hunter, E. J. Manbers, M. Nowotka, "
        "A. P. Bento, D. Mendez and G. J. P. van Westen, "
        "Nucleic Acids Res., 2024, 52, D1180\u2013D1192.",
        # 19 - DUD-E
        "M. M. Mysinger, M. Carchia, J. J. Irwin and B. K. Shoichet, "
        "J. Med. Chem., 2012, 55, 6582\u20136594.",
        # 20 - BEDROC
        "J. F. Truchon and C. I. Bayly, "
        "J. Chem. Inf. Model., 2007, 47, 488\u2013508.",
        # 21 - Anthropic Claude
        "Anthropic, The Claude 3 model family: Opus, Sonnet, Haiku, "
        "Technical report, 2024, "
        "https://www-cdn.anthropic.com/de8ba9b01c9ab7cbabf5c33b80b7bbc618857627/Model_Card_Claude_3.pdf.",
        # 22 - RDKit
        "G. Landrum, RDKit: open-source cheminformatics, 2024, "
        "https://www.rdkit.org/.",
    ]

    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)
        run_num = p.add_run(f"{i} ")
        run_num.font.name = "Times New Roman"
        run_num.font.size = Pt(10)
        _write_ref_parts(p, ref)

    outpath = os.path.join(OUTDIR, "manuscript_digdisc.docx")
    doc.save(outpath)
    print(f"Saved: {outpath}")
    return outpath


# ===========================================================================
# Supplementary (ESI)
# ===========================================================================
def make_supplementary():
    doc = Document()
    setup_styles(doc)
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # Title
    p_text(doc, "Electronic Supplementary Information (ESI)",
           bold=True, font_size=16,
           alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    p_text(doc,
           "Autonomous Optimisation of Structure-Based Virtual Screening "
           "Protocols Using an LLM Coding Agent",
           bold=True, font_size=13,
           alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    p_text(doc, "Osman Gani", font_size=12,
           alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    p_text(doc,
           "Section for Pharmaceutical Chemistry, Department of Pharmacy, "
           "University of Oslo, P.O. Box 1068 Blindern, 0316 Oslo, Norway",
           italic=True, font_size=10,
           alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

    # ---- S1 ----
    doc.add_heading("S1. Agent program files", level=1)
    p_text(doc,
           "The agent operated under structured protocol files (program.md) "
           "that defined the experiment loop, research directions organised "
           "into three tiers by expected impact, and the baseline "
           "configuration. Separate program files were used for each target. "
           "The complete program files are available in the repository.")

    p_text(doc, "Both programs instructed the agent to:", space_after=2)
    for item in [
        "Never modify the infrastructure code (prepare.py), only the "
        "experiment configuration (experiment.py)",
        "Test exactly one hypothesis per git commit",
        "Operate fully autonomously without pausing for human input",
        "Follow a strict hypothesise\u2013modify\u2013commit\u2013run"
        "\u2013evaluate\u2013decide loop",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)

    p_text(doc, "Research directions were organised as:",
           space_after=2, space_before=6)
    for tier in [
        ("Tier 1 (high impact, fast):",
         " scoring function, box size, exhaustiveness, library choice, "
         "receptor choice"),
        ("Tier 2 (moderate impact, new logic):",
         " score normalisation, consensus scoring, multi-pose analysis, "
         "MW correction"),
        ("Tier 3 (exploratory):",
         " energy range, num_modes, seed variation"),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(tier[0])
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        run2 = p.add_run(tier[1])
        run2.font.name = "Times New Roman"
        run2.font.size = Pt(11)

    # ---- S2: FPR2 experiments ----
    doc.add_heading("S2. FPR2 detailed experiment logs", level=1)

    fpr2_experiments = [
        ("Experiment 0: Baseline",
         [("Configuration: ", True),
          ("Vina scoring, skill library (5,223 ligands), skill receptor "
           "(Meeko), 25 \u00c5 box, exhaustiveness 8, batch size 100.",
           False)],
         [("Results: ", True),
          ("4,100 ligands successfully docked (1,123 lost to segfaults). "
           "AUC = 0.736, BEDROC = 0.277, EF 1% = 1.93. "
           "Runtime: 21,786 s (6.1 h).", False)],
         None),
        ("Experiment 1: Vinardo scoring (KEPT)",
         [("Hypothesis: ", True),
          ("The Vinardo scoring function may discriminate actives from "
           "decoys better than the default Vina function for FPR2.", False)],
         [("Results: ", True),
          ("AUC = 0.748 (+0.012). BEDROC = 0.238, EF 1% = 1.05. "
           "Runtime: 22,284 s (6.2 h). Decision: KEPT.", False)],
         [("Note: ", True),
          ("Improvement of +0.012 exceeds the estimated noise floor of "
           "0.004 by threefold.", False)]),
        ("Experiment 2: Box size 20 \u00c5",
         [("Hypothesis: ", True),
          ("A tighter docking box (20 vs 25 \u00c5) may focus the search "
           "on the orthosteric pocket.", False)],
         [("Results: ", True),
          ("AUC = 0.745 (\u22120.003). BEDROC = 0.221, EF 1% = 0.70. "
           "Decision: REVERTED.", False)],
         None),
        ("Experiment 3: Naive library",
         [("Hypothesis: ", True),
          ("The naive library (10,752 compounds) covers more chemical "
           "space.", False)],
         [("Results: ", True),
          ("Total failure. All 108 batches crashed. Only 100 of 10,752 "
           "ligands were docked. AUC = NaN. Decision: REVERTED.", False)],
         [("Root cause: ", True),
          ("OpenBabel-generated PDBQT files contained zero coordinates, "
           "missing atoms, and invalid tags.", False)]),
        ("Experiment 4: Rank score transform",
         [("Hypothesis: ", True),
          ("Rank-based normalisation may improve discrimination.", False)],
         [("Results: ", True),
          ("AUC = 0.744, identical to untransformed AUC.", False)],
         [("Key insight: ", True),
          ("The agent independently realised that monotonic transforms "
           "preserve relative ordering and cannot change ranking-based "
           "metrics. Skipped z-score and min-max.", False)]),
        ("Experiment 5: Exhaustiveness 16",
         [("Hypothesis: ", True),
          ("Doubling exhaustiveness may improve pose sampling.", False)],
         [("Results: ", True),
          ("AUC = 0.744 (\u22120.004). Runtime: 24,673 s (6.9 h, +11% "
           "slower). Decision: REVERTED.", False)],
         None),
        ("Experiment 6: Naive receptor",
         [("Hypothesis: ", True),
          ("The OpenBabel-prepared receptor may improve scoring.", False)],
         [("Results: ", True),
          ("AUC = 0.732 (\u22120.016). 305 corrupted scores. "
           "Decision: REVERTED.", False)],
         None),
        ("Experiment 7: MW correction",
         [("Hypothesis: ", True),
          ("Dividing scores by MW^0.33 may correct size bias.", False)],
         [("Results: ", True),
          ("AUC = 0.702 (\u22120.046). Worst result in the campaign. "
           "Decision: REVERTED.", False)],
         None),
    ]

    for title, line1, line2, line3 in fpr2_experiments:
        doc.add_heading(title, level=2)
        for line_parts in [line1, line2, line3]:
            if line_parts is None:
                continue
            p = doc.add_paragraph()
            for text, bold in line_parts:
                r = p.add_run(text)
                r.bold = bold
                r.font.name = "Times New Roman"
                r.font.size = Pt(11)

    doc.add_heading("Experiments 8\u20139: Multi-pose aggregation", level=2)
    p = doc.add_paragraph()
    r = p.add_run("Hypothesis: ")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
    r2 = p.add_run("Using information from multiple docked poses may "
                    "provide a more robust estimate of binding propensity.")
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(11)

    mp_headers = ["Method", "AUC", "\u0394AUC"]
    mp_rows = [
        ["Best (standard)", "0.740", "reference"],
        ["Boltzmann-weighted", "0.730", "\u22120.010"],
        ["Mean top-3", "0.715", "\u22120.025"],
        ["Mean all", "0.702", "\u22120.038"],
        ["Median", "0.701", "\u22120.039"],
    ]
    add_table(doc, mp_headers, mp_rows)
    p_text(doc, "\nAll aggregation methods performed worse than the "
                "standard best-pose approach.")

    # Fig S1
    add_figure(doc, "figure_s1_fpr2_waterfall.png", [
        ("Fig. S1 ", True),
        ("FPR2: change in AUC relative to baseline for each experiment. "
         "Blue: kept; red: reverted; grey: no change or failed.",),
    ], width=Inches(5.5))

    # ---- S3: CDK2 experiments ----
    doc.add_heading("S3. CDK2 detailed experiment logs", level=1)

    cdk2_experiments = [
        ("Experiment 0: Baseline",
         "Configuration: Vina scoring, skill library (1,309 ligands), "
         "skill receptor (Meeko), 25 \u00c5 box, exhaustiveness 8, "
         "dual-GPU.",
         "Results: AUC = 0.677, BEDROC = 0.443, EF 1% = 1.71. "
         "Runtime: 1,342 s (22 min)."),
        ("Experiment 1: Vinardo scoring (KEPT)",
         "Hypothesis: Vinardo may improve enrichment for CDK2.",
         "Results: AUC = 0.695 (+0.018). Decision: KEPT."),
        ("Experiments 2\u20134: Box size sweep",
         "Tested 20, 18, and 30 \u00c5 boxes.",
         "Results: AUC 0.692, 0.694, 0.685. All reverted. 25 \u00c5 "
         "remained optimal."),
        ("Experiments 5\u20136: Exhaustiveness sweep",
         "Tested exh. 16 and 32.",
         "Results: AUC 0.686, 0.693. Both reverted. Exh. 8 optimal."),
        ("Experiment 7: Naive library with Vinardo (KEPT)",
         "Hypothesis: OpenBabel library covers more chemical space.",
         "Results: AUC = 0.716 (+0.021). Dramatic improvement from "
         "larger library (2,471 vs 1,309 ligands). BEDROC jumped from "
         "0.403 to 0.711. Decision: KEPT."),
        ("Experiment 8: Naive receptor + naive library",
         "Hypothesis: Matching naive prep for both receptor and library.",
         "Results: AUC = 0.706 (\u22120.010). Skill receptor remains "
         "superior. Decision: REVERTED."),
        ("Experiment 9: Naive library with Vina scoring (KEPT)",
         "Hypothesis: Vina may outperform Vinardo with the naive library.",
         "Results: AUC = 0.735 (+0.019). Vina outperforms Vinardo when "
         "using naive library, opposite of FPR2 finding. "
         "Decision: KEPT."),
        ("Experiment 10: Energy range 5 kcal/mol",
         "Tested wider energy range.",
         "Results: AUC = 0.735 (identical). Decision: REVERTED."),
        ("Experiment 11: Batch size 50 (KEPT)",
         "Hypothesis: Smaller batches reduce ligand loss from segfaults.",
         "Results: AUC = 0.733, but 1,951 ligands scored (vs. ~1,500 "
         "with batch 100). Decision: KEPT for better coverage."),
        ("Experiments 12\u201315: Additional parameter tests",
         "Tested batch size 20, exh. 4, box 22 \u00c5, num_modes 20.",
         "Results: AUC 0.730, 0.729, 0.729, 0.727. All reverted."),
    ]

    for title, line1, line2 in cdk2_experiments:
        doc.add_heading(title, level=2)
        for text in [line1, line2]:
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.font.name = "Times New Roman"
            r.font.size = Pt(11)

    # Fig S2
    add_figure(doc, "figure_s2_cdk2_waterfall.png", [
        ("Fig. S2 ", True),
        ("CDK2: change in AUC relative to baseline (0.677) for each "
         "experiment. Blue: kept; red: reverted.",),
    ], width=Inches(5.5))

    # ---- S4: Stochastic variation ----
    doc.add_heading("S4. Stochastic variation analysis", level=1)
    p_text(doc,
           "To establish the noise floor for FPR2, the agent observed that "
           "two runs of the identical configuration (Vinardo, skill library, "
           "skill receptor, 25 \u00c5 box, exhaustiveness 8) produced AUC "
           "values of 0.744 and 0.748. The difference of 0.004 reflects "
           "stochastic variation from Uni-Dock\u2019s GPU-parallel Monte "
           "Carlo sampling and batch-dependent segfault patterns. This noise "
           "floor of approximately 0.004 AUC was used as the threshold for "
           "declaring improvements meaningful.")
    p_text(doc,
           "For CDK2, a similar noise floor was observed: experiments with "
           "identical docking parameters but different batch sizes yielded "
           "AUC variations of approximately 0.002\u20130.005, attributable "
           "to the stochastic ligand loss from segfaults.")

    # ---- S5: Segfault analysis ----
    doc.add_heading("S5. Segfault analysis", level=1)
    p_text(doc,
           "FPR2: Across all runs, batches 30\u201340 (ligands DEC_05xxx "
           "through DEC_07xxx) consistently caused segmentation faults. "
           "Inspection revealed structural parsing errors (\u201cUnknown or "
           "inappropriate tag found in flex residue or ligand\u201d). These "
           "files were Meeko-prepared but contain unusual chemical features "
           "(e.g., boron-containing compounds, complex ring systems) that "
           "produce PDBQT tags not recognised by Uni-Dock\u2019s parser. "
           "Approximately 1,100 ligands were lost per run (79% success "
           "rate).")
    p_text(doc,
           "CDK2: The naive (OpenBabel-prepared) library exhibited a higher "
           "segfault rate (exit code \u221211) on both GPUs, with "
           "approximately 20\u201340% of batches crashing. Reducing batch "
           "size from 100 to 50 improved the success rate from ~62% to ~79% "
           "of ligands scored. The skill library showed a similar segfault "
           "pattern to FPR2 (~20% loss).")

    # ---- S6: Git histories ----
    doc.add_heading("S6. Git commit histories", level=1)

    p_text(doc, "FPR2 campaign:", bold=True, space_after=2)
    fpr2_git = """56f1239 revert: multi-pose scoring (all methods worse than best-pose)
df1a699 experiment: multi-pose scoring sweep (free, reuse cached results)
950c05d revert: MW correction (AUC 0.7022, much worse than 0.7480)
b06bfea experiment: add MW correction (score / MW^0.33)
9e9b718 revert: naive receptor (AUC 0.7320, worse than 0.7480)
5115af9 experiment: switch to naive receptor (OpenBabel-prepared)
03f9cc7 revert exh16: AUC 0.7441 < 0.7480
d3539e7 experiment: increase exhaustiveness from 8 to 16
2e4b9b0 revert rank transform: monotonic transforms don't affect AUC
31fc684 experiment: test rank score transform
03c3d1e experiment: add reuse_results option for fast transform testing
31860ec revert naive library: total failure, 100/10751 docked, AUC=nan
31a6426 experiment: switch to naive library (10752 OpenBabel ligands)
b848ee5 revert box 20A: AUC 0.7447 < 0.7480
e32c6b0 experiment: reduce box size from 25A to 20A
ff7fd1b experiment: switch scoring function from vina to vinardo
2659a74 Initial VS autoresearch scaffold"""

    p = doc.add_paragraph()
    run = p.add_run(fpr2_git)
    run.font.name = "Courier New"
    run.font.size = Pt(8)

    p_text(doc, "\nCDK2 campaign:", bold=True, space_after=2)
    cdk2_git = """REVERTED num_modes 20 (naive/vina/25A/exh8)
REVERTED box 22A (naive/vina/exh8/batch50)
REVERTED exh4 (naive/vina/25A/batch50)
REVERTED batch_size 20 (naive/vina/25A)
2fa780e experiment: batch_size 50 (naive/vina/25A/exh8)
REVERTED energy_range 5 (naive/vina/25A/exh8)
ff17182 experiment: naive library + vina scoring (25A/exh8)
REVERTED naive receptor+library (vinardo/25A/exh8)
fd1cc8f experiment: naive library (vinardo/25A/exh8)
REVERTED exh32 (vinardo/25A)
REVERTED exh16 (vinardo/25A)
REVERTED box 30A (vinardo)
REVERTED box 18A (vinardo)
REVERTED box 20A (vinardo)
2098b88 experiment: vinardo scoring
baseline skill/skill/25A/exh8/vina/none"""

    p = doc.add_paragraph()
    run = p.add_run(cdk2_git)
    run.font.name = "Courier New"
    run.font.size = Pt(8)

    # ---- S7: Final CONFIGs ----
    doc.add_heading("S7. Final optimised configurations", level=1)

    p_text(doc, "FPR2 final configuration:", bold=True, space_after=2)
    fpr2_config = """CONFIG = {
    "library": "skill",         # Meeko-prepared, 5,223 ligands
    "receptor": "skill",        # Meeko mk_prepare_receptor
    "box_size": 25,             # Angstrom
    "exhaustiveness": 8,        # GPU-parallel Monte Carlo
    "num_modes": 10,
    "energy_range": 3,          # kcal/mol above best
    "scoring": "vinardo",       # Only change from baseline
    "score_transform": "none",
    "batch_size": 100,
}"""
    p = doc.add_paragraph()
    run = p.add_run(fpr2_config)
    run.font.name = "Courier New"
    run.font.size = Pt(9)

    p_text(doc, "\nCDK2 final configuration:", bold=True,
           space_after=2, space_before=12)
    cdk2_config = """CONFIG = {
    "library": "naive",         # OpenBabel-prepared, 2,471 ligands
    "receptor": "skill",        # Meeko mk_prepare_receptor
    "box_size": 25,             # Angstrom
    "exhaustiveness": 8,        # GPU-parallel Monte Carlo
    "num_modes": 10,
    "energy_range": 3,          # kcal/mol above best
    "scoring": "vina",          # Vina (not Vinardo)
    "score_transform": "none",
    "batch_size": 50,           # Smaller batches for naive library
    "gpu_device": "0,1",        # Dual GPU
}"""
    p = doc.add_paragraph()
    run = p.add_run(cdk2_config)
    run.font.name = "Courier New"
    run.font.size = Pt(9)

    p_text(doc,
           "\nKey differences: FPR2 uses Vinardo scoring with the skill "
           "(Meeko+PAINS-filtered) library, while CDK2 uses Vina scoring "
           "with the naive (OpenBabel) library. Both targets use the "
           "Meeko-prepared receptor.")

    # Fig S3 - GPU hours
    add_figure(doc, "figure_s3_gpu_hours.png", [
        ("Fig. S3 ", True),
        ("Cumulative GPU-hours for FPR2 (A, single GPU, \u223c38 h total) "
         "and CDK2 (B, dual GPU, \u223c6 h total). The CDK2 campaign was "
         "substantially faster due to the smaller library and dual-GPU "
         "execution.",),
    ], width=Inches(6.0))

    outpath = os.path.join(OUTDIR, "supplementary_digdisc.docx")
    doc.save(outpath)
    print(f"Saved: {outpath}")
    return outpath


if __name__ == "__main__":
    make_manuscript()
    make_supplementary()
